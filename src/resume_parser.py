import io
import re

import docx
import pdfplumber

NOT_CAPTURED = "(not captured)"

_CANDIDATE_RE = re.compile(r"\+?\d[\d\s\-]{7,15}\d")
_FALLBACK_INTL_RE = re.compile(r"\+\d{7,15}")
_LABEL_RE = re.compile(r"(phone|mobile|contact|cell|tel)", re.IGNORECASE)
_CONTEXT_CHARS = 15


def extract_text(filename, data_bytes):
    lowered = filename.lower()
    if lowered.endswith(".pdf"):
        return _extract_pdf_text(data_bytes)
    if lowered.endswith(".docx"):
        return _extract_docx_text(data_bytes)
    return ""


def _extract_pdf_text(data_bytes):
    chunks = []
    with pdfplumber.open(io.BytesIO(data_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                chunks.append(text)
    return "\n".join(chunks)


def _extract_docx_text(data_bytes):
    document = docx.Document(io.BytesIO(data_bytes))
    chunks = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    chunks.append(cell.text)
    return "\n".join(chunks)


def _looks_like_year_run(match_text):
    groups = [g for g in re.split(r"\D+", match_text) if g]
    year_like = [g for g in groups if len(g) == 4 and 1900 <= int(g) <= 2099]
    return len(year_like) >= 2


def _score_match(text, match):
    digits = re.sub(r"\D", "", match.group())
    digit_count = len(digits)
    if digit_count < 10 or digit_count > 13:
        return None
    if _looks_like_year_run(match.group()):
        return None

    score = 0
    if digit_count == 10:
        score += 3
    elif digit_count == 12 and digits.startswith("91"):
        score += 3
    else:
        score += 1

    if match.group().strip().startswith("+91"):
        score += 2

    context_before = text[max(0, match.start() - _CONTEXT_CHARS):match.start()]
    if _LABEL_RE.search(context_before):
        score += 2

    before_char = text[match.start() - 1] if match.start() > 0 else ""
    after_char = text[match.end()] if match.end() < len(text) else ""
    if before_char == "/" or after_char == "/":
        score -= 3

    return score


def extract_phone(text):
    if not text:
        return NOT_CAPTURED

    best_score = None
    best_match = None
    for match in _CANDIDATE_RE.finditer(text):
        score = _score_match(text, match)
        if score is None:
            continue
        if best_score is None or score > best_score:
            best_score = score
            best_match = match

    if best_match is not None and best_score >= 1:
        return best_match.group().strip()

    fallback = _FALLBACK_INTL_RE.search(text)
    if fallback:
        return fallback.group().strip()

    return NOT_CAPTURED
